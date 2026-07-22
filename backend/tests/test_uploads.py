import io
from pathlib import Path

from docx import Document
from fastapi.testclient import TestClient
from PIL import Image
from reportlab.pdfgen import canvas
from sqlalchemy.orm import Session, sessionmaker

from app.core.config import settings
from app.models.book import Book


def _create_book(client: TestClient) -> dict[str, object]:
    response = client.post(
        "/api/v1/books",
        json={
            "title": "Upload Test Book",
            "owner_name": "Alex",
            "number_mode": "none",
        },
    )
    assert response.status_code == 201
    return response.json()


def _docx_bytes() -> bytes:
    output = io.BytesIO()
    document = Document()
    document.add_heading("Uploaded section", level=1)
    document.add_paragraph("中英文 DOCX content")
    document.save(output)
    return output.getvalue()


def _pdf_bytes() -> bytes:
    output = io.BytesIO()
    pdf = canvas.Canvas(output)
    pdf.drawString(72, 720, "Uploaded PDF section")
    pdf.showPage()
    pdf.save()
    return output.getvalue()


def _image_bytes(image_format: str) -> bytes:
    output = io.BytesIO()
    Image.new("RGB", (40, 30), "#4f7cff").save(output, format=image_format)
    return output.getvalue()


def test_upload_preview_replace_and_delete_use_one_storage_chain(
    client: TestClient,
    monkeypatch,
    tmp_path: Path,
    test_session_factory: sessionmaker[Session],
) -> None:
    monkeypatch.setattr(settings, "storage_dir", tmp_path)
    book = _create_book(client)
    book_id = book["id"]

    uploads = {
        "cover": ("cover.png", _image_bytes("PNG"), "image/png"),
        "preface": ("preface.pdf", _pdf_bytes(), "application/pdf"),
        "afterword": (
            "afterword.docx",
            _docx_bytes(),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ),
        "acknowledgement": (
            "thanks.jpeg",
            _image_bytes("JPEG"),
            "image/jpeg",
        ),
        "back_cover": (
            "back.webp",
            _image_bytes("WEBP"),
            "image/webp",
        ),
    }

    for upload_type, (name, content, media_type) in uploads.items():
        response = client.post(
            f"/api/v1/books/{book_id}/upload",
            data={"type": upload_type},
            files={"file": (name, content, media_type)},
        )
        assert response.status_code == 200
        payload = response.json()
        expected_path = f"books/{book_id}/{upload_type}/{name}"
        assert payload["success"] is True
        assert payload["file_name"] == name
        assert payload["file_size"] == len(content)
        assert payload["file_type"] == media_type
        assert payload["path"] == expected_path
        assert payload["uploaded_at"]
        assert (tmp_path / expected_path).read_bytes() == content

    with test_session_factory() as session:
        persisted = session.get(Book, book_id)
        assert persisted is not None
        assert persisted.cover_file == f"books/{book_id}/cover/cover.png"
        assert persisted.preface_file == f"books/{book_id}/preface/preface.pdf"
        assert persisted.afterword_file == f"books/{book_id}/afterword/afterword.docx"
        assert persisted.acknowledgement_file == (
            f"books/{book_id}/acknowledgement/thanks.jpeg"
        )
        assert persisted.back_cover_file == f"books/{book_id}/back_cover/back.webp"

    preview = client.get(f"/api/v1/files/{book_id}/preface")
    assert preview.status_code == 200
    assert preview.content == uploads["preface"][1]
    assert preview.headers["content-type"] == "application/pdf"
    assert preview.headers["content-disposition"].startswith("inline")

    metadata = client.head(f"/api/v1/files/{book_id}/preface")
    assert metadata.status_code == 200
    assert int(metadata.headers["content-length"]) == len(uploads["preface"][1])
    assert metadata.headers["x-file-path"] == (f"books/{book_id}/preface/preface.pdf")
    assert metadata.headers["last-modified"]

    replacement = _docx_bytes()
    replaced = client.post(
        f"/api/v1/books/{book_id}/upload",
        data={"type": "preface"},
        files={
            "file": (
                "latest-preface.docx",
                replacement,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert replaced.status_code == 200
    assert not (tmp_path / f"books/{book_id}/preface/preface.pdf").exists()
    assert (
        tmp_path / f"books/{book_id}/preface/latest-preface.docx"
    ).read_bytes() == replacement

    deleted = client.delete(f"/api/v1/books/{book_id}/upload/preface")
    assert deleted.status_code == 204
    assert client.get(f"/api/v1/files/{book_id}/preface").status_code == 404
    assert not (tmp_path / f"books/{book_id}/preface/latest-preface.docx").exists()
    with test_session_factory() as session:
        persisted = session.get(Book, book_id)
        assert persisted is not None
        assert persisted.preface_file is None


def test_upload_updates_matching_layout_section(
    client: TestClient,
    monkeypatch,
    tmp_path: Path,
) -> None:
    monkeypatch.setattr(settings, "storage_dir", tmp_path)
    book = _create_book(client)
    book_id = book["id"]
    client.patch(
        f"/api/v1/books/{book_id}",
        json={
            "layout_sections": [
                {
                    "id": "cover",
                    "kind": "page",
                    "preset": "cover",
                    "name": None,
                    "file": None,
                },
                {
                    "id": "articles",
                    "kind": "articles",
                    "preset": "articles",
                    "name": None,
                    "file": None,
                },
                {
                    "id": "back_cover",
                    "kind": "page",
                    "preset": "back_cover",
                    "name": None,
                    "file": None,
                },
            ]
        },
    )

    uploaded = client.post(
        f"/api/v1/books/{book_id}/upload",
        data={"type": "cover"},
        files={"file": ("cover.jpg", _image_bytes("JPEG"), "image/jpeg")},
    )
    assert uploaded.status_code == 200
    refreshed = client.get(f"/api/v1/books/{book_id}").json()
    assert refreshed["layout_sections"][0]["file"] == (
        f"books/{book_id}/cover/cover.jpg"
    )

    protected = client.delete(f"/api/v1/books/{book_id}/upload/cover")
    assert protected.status_code == 409
    assert protected.json()["detail"]["code"] == "protected_upload"
    refreshed = client.get(f"/api/v1/books/{book_id}").json()
    assert refreshed["layout_sections"][0]["file"] == (
        f"books/{book_id}/cover/cover.jpg"
    )


def test_upload_validation_errors_are_bilingual_and_leave_no_file(
    client: TestClient,
    monkeypatch,
    tmp_path: Path,
) -> None:
    monkeypatch.setattr(settings, "storage_dir", tmp_path)
    book_id = _create_book(client)["id"]

    unsupported = client.post(
        f"/api/v1/books/{book_id}/upload",
        data={"type": "preface"},
        files={"file": ("preface.txt", b"plain text", "text/plain")},
    )
    assert unsupported.status_code == 415
    assert unsupported.json()["detail"] == {
        "code": "unsupported_file_format",
        "message": "Unsupported file format.",
        "message_zh": "不支持的文件格式。",
    }

    fake_pdf = client.post(
        f"/api/v1/books/{book_id}/upload",
        data={"type": "preface"},
        files={"file": ("preface.pdf", b"not a pdf", "application/pdf")},
    )
    assert fake_pdf.status_code == 415

    monkeypatch.setattr(settings, "max_upload_size", 8)
    too_large = client.post(
        f"/api/v1/books/{book_id}/upload",
        data={"type": "preface"},
        files={"file": ("preface.pdf", b"%PDF-1.7-too-large", "application/pdf")},
    )
    assert too_large.status_code == 413
    assert too_large.json()["detail"]["message_zh"] == "文件超过 100 MB 限制。"

    assert client.get(f"/api/v1/files/{book_id}/preface").status_code == 404
    assert client.delete("/api/v1/books/999/upload/cover").status_code == 404
    assert (
        client.post(
            "/api/v1/books/999/upload",
            data={"type": "cover"},
            files={"file": ("cover.png", _image_bytes("PNG"), "image/png")},
        ).status_code
        == 404
    )
