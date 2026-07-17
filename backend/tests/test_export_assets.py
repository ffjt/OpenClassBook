import io
from pathlib import Path

import pytest
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from fastapi.testclient import TestClient
from PIL import Image
from pypdf import PdfReader, PdfWriter
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas

from app.core.config import settings
from app.schemas.export import ExportTemplateInfo
from app.services.docx_formatting import convert_docx_to_html
from app.services.docx_word_converter import WordConversionResult
from app.services.page_asset_renderer import PageAssetRenderer


@pytest.fixture(autouse=True)
def _disable_native_word_conversion(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr(
        "app.services.page_asset_renderer.convert_docx_with_word",
        lambda source, destination: WordConversionResult(
            converted=False,
            attempted=False,
        ),
    )


def _image_bytes(image_format: str = "PNG") -> bytes:
    output = io.BytesIO()
    image = Image.new("RGBA", (180, 120), (79, 124, 255, 180))
    if image_format in {"JPEG", "WEBP"}:
        image = image.convert("RGB")
    image.save(output, format=image_format)
    return output.getvalue()


def _docx_bytes() -> bytes:
    output = io.BytesIO()
    document = Document()
    document.add_heading("DOCX section heading", level=1)
    document.add_paragraph("DOCX section content / 文档内容")
    document.add_paragraph("First list item", style="List Bullet")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Language"
    table.cell(1, 1).text = "中英文"
    document.add_picture(io.BytesIO(_image_bytes()), width=Inches(1.2))
    document.save(output)
    return output.getvalue()


def _rich_docx_bytes() -> bytes:
    output = io.BytesIO()
    document = Document()
    document.add_heading("Theme heading / 主题标题", level=1)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(18)
    paragraph.paragraph_format.first_line_indent = Pt(18)
    run = paragraph.add_run("Colored DOCX marker / 彩色文字")
    run.bold = True
    run.underline = True
    run.font.color.rgb = RGBColor(255, 0, 0)
    run.font.highlight_color = 7
    run.font.size = Pt(18)
    document.add_picture(io.BytesIO(_image_bytes()), width=Inches(2))
    cell = document.add_table(rows=1, cols=1).cell(0, 0)
    cell.text = "Shaded table cell"
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "00FF00")
    cell._tc.get_or_add_tcPr().append(shading)
    document.add_page_break()
    document.add_paragraph("Second DOCX page")
    document.save(output)
    return output.getvalue()


def _export_template() -> ExportTemplateInfo:
    return ExportTemplateInfo(
        font="noto-serif-sc",
        font_size=14,
        page_size="a4",
        page_margin="normal",
        allow_images=True,
        image_align="center",
        image_width=80,
        numbering_style="arabic",
        line_height=1.5,
        title_size=24,
        title_align="center",
        title_bold=True,
        subtitle_mode="free",
        fixed_subtitle="",
        subtitle_align="center",
        body_justify=True,
        first_line_indent=2,
        page_number_position="center",
        custom_page_width=210,
        custom_page_height=297,
    )


def _pdf_bytes(*labels: str) -> bytes:
    output = io.BytesIO()
    pdf = canvas.Canvas(output, pagesize=(360, 540))
    for label in labels:
        pdf.drawString(48, 480, label)
        pdf.showPage()
    pdf.save()
    return output.getvalue()


def _encrypted_pdf_bytes() -> bytes:
    source = PdfReader(io.BytesIO(_pdf_bytes("Secret page")))
    writer = PdfWriter()
    writer.add_page(source.pages[0])
    writer.encrypt("secret")
    output = io.BytesIO()
    writer.write(output)
    return output.getvalue()


def _create_book(client: TestClient, title: str = "Parsed assets") -> dict[str, object]:
    return client.post(
        "/api/v1/books",
        json={"title": title, "owner_name": "Editor", "number_mode": "none"},
    ).json()


def _upload(
    client: TestClient,
    book_id: int,
    upload_type: str,
    name: str,
    content: bytes,
    media_type: str,
) -> None:
    response = client.post(
        f"/api/v1/books/{book_id}/upload",
        data={"type": upload_type},
        files={"file": (name, content, media_type)},
    )
    assert response.status_code == 200, response.text


def test_mixed_docx_pdf_and_images_are_rendered_in_layout_order(
    client: TestClient,
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
) -> None:
    storage_dir = tmp_path / "storage"
    export_dir = tmp_path / "exports"
    monkeypatch.setattr(settings, "storage_dir", storage_dir)
    monkeypatch.setattr(settings, "export_dir", export_dir)
    book = _create_book(client)
    book_id = int(book["id"])
    client.patch(
        f"/api/v1/books/{book_id}",
        json={
            "layout_sections": [
                {"id": "cover", "kind": "page", "preset": "cover"},
                {"id": "preface", "kind": "page", "preset": "preface"},
                {"id": "articles", "kind": "articles", "preset": "articles"},
                {"id": "afterword", "kind": "page", "preset": "afterword"},
                {
                    "id": "back_cover",
                    "kind": "page",
                    "preset": "back_cover",
                },
            ]
        },
    )
    _upload(client, book_id, "cover", "cover.png", _image_bytes(), "image/png")
    _upload(
        client,
        book_id,
        "preface",
        "preface.docx",
        _docx_bytes(),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    _upload(
        client,
        book_id,
        "afterword",
        "afterword.pdf",
        _pdf_bytes("PDF afterword page one", "PDF afterword page two"),
        "application/pdf",
    )
    _upload(
        client,
        book_id,
        "back_cover",
        "back.webp",
        _image_bytes("WEBP"),
        "image/webp",
    )
    author = client.post(
        f"/api/v1/books/{book_id}/authors",
        json={"name": "Student"},
    ).json()
    article = client.post(
        f"/api/v1/books/{book_id}/articles",
        json={
            "author_id": author["id"],
            "title": "Approved article marker",
            "content": "Approved article content",
            "status": "approved",
        },
    )
    assert article.status_code == 201

    preview = client.get(f"/api/v1/books/{book_id}/export").json()
    assert preview["can_export"] is True
    assert preview["stats"]["estimated_page_count"] == 6
    assert preview["stats"]["image_count"] == 3
    assert [page["label_en"] for page in preview["preview_pages"]] == [
        "Cover",
        "Preface",
        "Approved article marker",
        "Afterword · 1/2",
        "Afterword · 2/2",
        "Back cover",
    ]
    assert [page["is_placeholder"] for page in preview["preview_pages"]] == [
        False,
        False,
        True,
        False,
        False,
        False,
    ]

    generated = client.post(f"/api/v1/books/{book_id}/export")
    assert generated.status_code == 200, generated.text
    result = generated.json()
    assert result["page_count"] == 6
    artifact = export_dir / f"book-{book_id}-{result['task_id']}.pdf"
    reader = PdfReader(artifact)
    assert len(reader.pages) == 6
    assert all(
        abs(float(page.mediabox.width) - A4[0]) < 0.1
        and abs(float(page.mediabox.height) - A4[1]) < 0.1
        for page in reader.pages
    )
    text = "\n".join(page.extract_text() or "" for page in reader.pages)
    assert text.index("DOCX section content") < text.index("Approved article marker")
    assert text.index("Approved article marker") < text.index("PDF afterword page one")
    assert text.index("PDF afterword page one") < text.index("PDF afterword page two")


def test_docx_direct_formatting_and_page_breaks_are_preserved(tmp_path: Path) -> None:
    source = tmp_path / "rich-formatting.docx"
    destination = tmp_path / "rich-formatting.pdf"
    source.write_bytes(_rich_docx_bytes())

    converted = convert_docx_to_html(source)
    assert "docx-run-" in converted.value
    assert "<strong><u><span" in converted.value
    assert "Colored DOCX marker / " in converted.value
    assert "彩色文字" in converted.value
    assert "docx-font-serif" in converted.value
    assert "docx-font-cjk" in converted.value
    assert "color: #365f91" in converted.css
    assert "color: #ff0000" in converted.css
    assert "background-color: #ffff00" in converted.css
    assert "font-size: 18pt" in converted.css
    assert "text-align: center" in converted.css
    assert "text-indent: 18pt" in converted.css
    assert "width: 144pt" in converted.css
    assert "background-color: #00ff00" in converted.css
    assert "display: block; page-break-after: always" in converted.css

    rendered = PageAssetRenderer().render(source, _export_template(), destination)
    assert rendered.page_count == 2
    assert rendered.image_count == 1
    reader = PdfReader(destination)
    first_page = reader.pages[0].extract_text() or ""
    assert "Colored DOCX marker" in first_page
    assert "彩色文字" in first_page
    assert "Second DOCX page" in (reader.pages[1].extract_text() or "")
    content = b"\n".join(page.get_contents().get_data() for page in reader.pages)
    assert b"1 0 0 rg" in content
    assert b"0 1 0 rg" in content


def test_docx_bullets_use_a_pdf_safe_marker_font(tmp_path: Path) -> None:
    source = tmp_path / "bullets.docx"
    destination = tmp_path / "bullets.pdf"
    source.write_bytes(_docx_bytes())
    converted = convert_docx_to_html(source)

    rendered = PageAssetRenderer().render(source, _export_template(), destination)

    assert "<ul><li" in converted.value
    assert rendered.page_count >= 1
    reader = PdfReader(destination)
    text = "\n".join(page.extract_text() or "" for page in reader.pages)
    assert "First list item" in text
    content = b"\n".join(page.get_contents().get_data() for page in reader.pages)
    assert b"/F1 14 Tf" in content
    assert b"(\\177) Tj" in content


def test_installed_word_is_preferred_over_compatible_renderer(
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
) -> None:
    source = tmp_path / "word-native.docx"
    destination = tmp_path / "word-native.pdf"
    source.write_bytes(_rich_docx_bytes())

    def convert_with_word(source_path: Path, destination_path: Path):
        assert source_path == source
        destination_path.write_bytes(_pdf_bytes("Native Microsoft Word output"))
        return WordConversionResult(converted=True, attempted=True)

    monkeypatch.setattr(
        "app.services.page_asset_renderer.convert_docx_with_word",
        convert_with_word,
    )
    monkeypatch.setattr(
        "app.services.page_asset_renderer.convert_docx_to_html",
        lambda source_path: pytest.fail(
            f"Compatible renderer should not run for {source_path}"
        ),
    )

    rendered = PageAssetRenderer().render(source, _export_template(), destination)
    assert rendered.page_count == 1
    assert rendered.image_count == 1
    assert rendered.warnings == ()
    assert "Native Microsoft Word output" in (
        PdfReader(destination).pages[0].extract_text() or ""
    )


def test_failed_word_conversion_falls_back_with_bilingual_warning(
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
) -> None:
    source = tmp_path / "word-fallback.docx"
    destination = tmp_path / "word-fallback.pdf"
    source.write_bytes(_rich_docx_bytes())
    monkeypatch.setattr(
        "app.services.page_asset_renderer.convert_docx_with_word",
        lambda source_path, destination_path: WordConversionResult(
            converted=False,
            attempted=True,
        ),
    )

    rendered = PageAssetRenderer().render(source, _export_template(), destination)
    assert rendered.page_count == 2
    assert rendered.warnings == (
        "Microsoft Word conversion failed; the compatible DOCX renderer was used.",
    )
    assert rendered.warnings_zh == (
        "Microsoft Word 转换失败，已改用兼容 DOCX 渲染器。",
    )


def test_rich_docx_upload_and_export_preserve_formatting_end_to_end(
    client: TestClient,
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
) -> None:
    storage_dir = tmp_path / "storage"
    export_dir = tmp_path / "exports"
    monkeypatch.setattr(settings, "storage_dir", storage_dir)
    monkeypatch.setattr(settings, "export_dir", export_dir)
    book_id = int(_create_book(client, "Rich DOCX export")["id"])
    response = client.patch(
        f"/api/v1/books/{book_id}",
        json={
            "layout_sections": [
                {"id": "preface", "kind": "page", "preset": "preface"},
                {"id": "articles", "kind": "articles", "preset": "articles"},
            ]
        },
    )
    assert response.status_code == 200, response.text
    _upload(
        client,
        book_id,
        "preface",
        "rich-formatting.docx",
        _rich_docx_bytes(),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    preview = client.get(f"/api/v1/books/{book_id}/export")
    assert preview.status_code == 200, preview.text
    preview_payload = preview.json()
    assert preview_payload["can_export"] is True
    assert preview_payload["stats"]["estimated_page_count"] == 2
    assert [page["label_en"] for page in preview_payload["preview_pages"]] == [
        "Preface · 1/2",
        "Preface · 2/2",
    ]

    generated = client.post(f"/api/v1/books/{book_id}/export")
    assert generated.status_code == 200, generated.text
    result = generated.json()
    assert result["page_count"] == 2
    artifact = export_dir / f"book-{book_id}-{result['task_id']}.pdf"
    reader = PdfReader(artifact)
    assert "Colored DOCX marker" in (reader.pages[0].extract_text() or "")
    assert "彩色文字" in (reader.pages[0].extract_text() or "")
    assert "Second DOCX page" in (reader.pages[1].extract_text() or "")
    content = b"\n".join(page.get_contents().get_data() for page in reader.pages)
    assert b"1 0 0 rg" in content
    assert b"0 1 0 rg" in content


def test_export_preview_skips_native_word_conversion(
    client: TestClient,
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
) -> None:
    monkeypatch.setattr(settings, "storage_dir", tmp_path / "storage")
    monkeypatch.setattr(settings, "export_dir", tmp_path / "exports")
    book_id = int(_create_book(client, "Fast DOCX preview")["id"])
    _upload(
        client,
        book_id,
        "preface",
        "preface.docx",
        _docx_bytes(),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    native_calls = 0

    def count_native_calls(source: Path, destination: Path) -> WordConversionResult:
        nonlocal native_calls
        native_calls += 1
        return WordConversionResult(converted=False, attempted=True)

    monkeypatch.setattr(
        "app.services.page_asset_renderer.convert_docx_with_word",
        count_native_calls,
    )

    preview = client.get(f"/api/v1/books/{book_id}/export")

    assert preview.status_code == 200, preview.text
    assert preview.json()["can_export"] is True
    assert native_calls == 0


def test_lightweight_export_preview_does_not_render_uploaded_assets(
    client: TestClient,
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
) -> None:
    monkeypatch.setattr(settings, "storage_dir", tmp_path / "storage")
    monkeypatch.setattr(settings, "export_dir", tmp_path / "exports")
    book_id = int(_create_book(client, "On-demand preview")["id"])
    _upload(
        client,
        book_id,
        "preface",
        "preface.docx",
        _docx_bytes(),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    def fail_if_rendered(*args: object, **kwargs: object) -> None:
        raise AssertionError("lightweight preview must not render assets")

    monkeypatch.setattr(PageAssetRenderer, "render", fail_if_rendered)

    preview = client.get(f"/api/v1/books/{book_id}/export?preflight=false")

    assert preview.status_code == 200, preview.text
    payload = preview.json()
    assert payload["can_export"] is True
    preface = next(
        page for page in payload["preview_pages"] if page["label_en"] == "Preface"
    )
    assert preface["is_placeholder"] is False


def test_uploaded_asset_without_approved_articles_can_export(
    client: TestClient,
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
) -> None:
    monkeypatch.setattr(settings, "storage_dir", tmp_path / "storage")
    monkeypatch.setattr(settings, "export_dir", tmp_path / "exports")
    book = _create_book(client, "Assets only")
    book_id = int(book["id"])
    _upload(
        client,
        book_id,
        "preface",
        "preface.pdf",
        _pdf_bytes("Only uploaded content"),
        "application/pdf",
    )

    preview = client.get(f"/api/v1/books/{book_id}/export").json()
    assert preview["stats"]["article_count"] == 0
    assert preview["can_export"] is True
    generated = client.post(f"/api/v1/books/{book_id}/export")
    assert generated.status_code == 200
    assert generated.json()["page_count"] >= 1


def test_missing_source_blocks_preview_and_returns_bilingual_export_error(
    client: TestClient,
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
) -> None:
    storage_dir = tmp_path / "storage"
    monkeypatch.setattr(settings, "storage_dir", storage_dir)
    monkeypatch.setattr(settings, "export_dir", tmp_path / "exports")
    book = _create_book(client, "Missing source")
    book_id = int(book["id"])
    _upload(
        client,
        book_id,
        "preface",
        "preface.pdf",
        _pdf_bytes("Will be removed"),
        "application/pdf",
    )
    (storage_dir / f"books/{book_id}/preface/preface.pdf").unlink()

    preview = client.get(f"/api/v1/books/{book_id}/export").json()
    assert preview["can_export"] is False
    assert "前言文件不存在，请替换后再导出。" in preview["warnings_zh"]
    generated = client.post(f"/api/v1/books/{book_id}/export")
    assert generated.status_code == 422
    assert generated.json()["detail"] == {
        "code": "source_file_unreadable",
        "message": "Preface file cannot be parsed. Replace it before export.",
        "message_zh": "前言文件无法解析，请替换后再导出。",
    }


def test_encrypted_pdf_is_rejected_during_upload(
    client: TestClient,
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
) -> None:
    monkeypatch.setattr(settings, "storage_dir", tmp_path)
    book_id = int(_create_book(client, "Encrypted source")["id"])
    response = client.post(
        f"/api/v1/books/{book_id}/upload",
        data={"type": "preface"},
        files={
            "file": (
                "protected.pdf",
                _encrypted_pdf_bytes(),
                "application/pdf",
            )
        },
    )
    assert response.status_code == 415
    assert response.json()["detail"]["message_zh"] == "不支持的文件格式。"
